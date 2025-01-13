import os
import streamlit as st
import re
from docx import Document
import io
from dotenv import load_dotenv
import google.generativeai as genai
import PyPDF2
from docx.shared import Pt

# Load environment variables from .env
load_dotenv()

# Configure Gemini API
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Predefined list of skills for regex matching
PREDEFINED_SKILLS = [
    "Python", "Java", "SQL", "Machine Learning", "Data Analysis",
    "Data Engineering", "AWS", "Azure", "Docker", "Kubernetes",
    "ETL", "Big Data", "Hadoop", "Spark", "Tableau", "Power BI",
    "AWS Glue", "PySpark", "Aurora DB", "Dynamo DB", "Redshift",
    "Data Warehousing", "CI/CD", "Stone branch", "Scheduling Tool"
]

# Function to extract multi-word skills using regex
def match_skill(text, skills_list):
    matched_skills = set()
    for skill in skills_list:
        if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
            matched_skills.add(skill)
    return matched_skills

# Function to extract skills using Gemini API
def gemini_extract_skills(text):
    try:
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }
        model = genai.GenerativeModel(model_name="gemini-1.5-flash", generation_config=generation_config)
        chat_session = model.start_chat()
        prompt = (
            "Extract all skills from the following job description text. "
            "Ensure the response contains only skill names, separated by commas:\n\n"
            f"{text}"
        )
        response = chat_session.send_message(prompt)
        skills_text = response.text.strip()
        skills = {skill.strip() for skill in skills_text.split(",") if skill.strip()}
        return skills
    except Exception as e:
        st.error(f"Failed to extract skills using Gemini API: {e}")
        return set()

# Extract text from Word
def extract_text_from_docx(docx_file):
    text = ''
    try:
        doc = Document(docx_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
    return text

# Extract text from PDF
def extract_text_from_pdf(pdf_file):
    text = ''
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return text

# Combine skills from regex and Gemini
def extract_skills(job_description):
    regex_skills = match_skill(job_description, PREDEFINED_SKILLS)
    gemini_skills = gemini_extract_skills(job_description)
    return list(regex_skills.union(gemini_skills))

# Generate interview questions using Gemini
def generate_interview_questions(skills, experience_level, complexity, num_questions):
    try:
        model = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config={
                "temperature": 1,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": 8192,
                "response_mime_type": "text/plain",
            },
        )
        chat_session = model.start_chat()
        prompt = (
            f"Generate {num_questions} interview questions and answers for a professional "
            f"with skills: {', '.join(skills)}, "
            f"{experience_level} years of experience. Questions should be of {complexity} complexity."
        )
        response = chat_session.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error fetching questions: {e}")
        return ""

# Export to Word with proper formatting and alignment
def export_to_word(data, job_name):
    try:
        doc = Document()
        doc.add_heading(f'{job_name} - Interview Questions and Answers', 0)
        
        # For each Q&A pair, we ensure the formatting is uniform
        for qa in data:
            # Add Question with a uniform format
            question_paragraph = doc.add_paragraph(qa['Question'])
            question_paragraph.alignment = 1  # Center-align the question
            
            # Add Answer with a uniform format
            answer_paragraph = doc.add_paragraph(qa['Answer'])
            answer_paragraph.alignment = 0  # Left-align the answer
            
            # Optional: Adjust font size for readability
            for para in doc.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(12)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error exporting to Word: {e}")
        return None

# Main Streamlit app
def main():
    st.set_page_config(page_title="Mazo", page_icon="📄")
    # Centered logo
    st.markdown("""
        <style>
            .center-image {
                display: flex;
                justify-content: center;
                align-items: center;
                margin-bottom: 20px;
            }
        </style>
        <div class="center-image">
            <img src="https://mazobeam.com/wp-content/uploads/2023/12/mazoid-1.png" alt="MazoBot Logo" width="200"/>
        </div>
    """, unsafe_allow_html=True)
    st.title("Job Description Skills Extractor & Interview Question Generator")
    
    # Toggle between uploading JD or manually entering skills
    input_method = st.radio("Select Input Method:", ["Upload Job Description", "Manually Input Skills"])

    skills = []
    job_description = ""

    if input_method == "Upload Job Description":
        uploaded_file = st.file_uploader("Upload Job Description (Word or PDF)", type=["docx", "pdf"])

        if uploaded_file:
            job_name = os.path.splitext(uploaded_file.name)[0]
            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                job_description = extract_text_from_docx(uploaded_file)
            elif uploaded_file.type == "application/pdf":
                job_description = extract_text_from_pdf(uploaded_file)
            else:
                st.error("Unsupported file type.")
                return

            if job_description.strip():
                skills = extract_skills(job_description)
                if skills:
                    st.success("Skills extracted successfully!")
                    st.write("### Extracted Skills:")
                    st.write(", ".join(skills))
                else:
                    st.warning("No skills found in the job description.")
            else:
                st.error("Failed to extract text from the uploaded document.")
    
    elif input_method == "Manually Input Skills":
        st.info("Manually input skills below, separated by commas.")
        skills_input = st.text_area("Enter Skills", placeholder="e.g., Python, SQL, Data Analysis")
        if skills_input.strip():
            skills = [skill.strip() for skill in skills_input.split(",") if skill.strip()]
            st.write("### Entered Skills:")
            st.write(", ".join(skills))

    if skills:
        experience_level = st.number_input("Experience Level (years)", min_value=1, max_value=50, step=1, value=10)
        complexity = st.radio("Select Question Complexity", ["Basic", "Intermediate", "Advanced"])
        num_questions = st.number_input("Number of Questions to Generate", min_value=1, max_value=100, step=1, value=10)

        if st.button("Generate Questions"):
            st.info("Generating interview questions. Please wait...")
            generated_content = generate_interview_questions(skills, experience_level, complexity, num_questions)

            if generated_content:
                st.success("Questions generated successfully!")
                st.write("Generated Questions and Answers")
                st.text_area("Questions & Answers", generated_content, height=300)
                qa_pairs = [
                    {"Question": q.strip(), "Answer": a.strip()}
                    for q, a in zip(*[iter(generated_content.splitlines())] * 2)
                ]

                # Export to Word
                word_file = export_to_word(qa_pairs, "Job_Description" if input_method == "Upload Job Description" else "Manual_Skills")
                if word_file:
                    st.download_button(
                        label="Download as Word",
                        data=word_file,
                        file_name=f"{'Job_Description' if input_method == 'Upload Job Description' else 'Manual_Skills'}_Interview_Questions_&_Answers.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    main()
